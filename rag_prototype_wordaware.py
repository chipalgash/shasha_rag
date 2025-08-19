#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
RAG prototype for Russian normative DOCX documents with:
- Word-aware & table-aware chunking (sentences/clauses, no mid-word breaks)
- Hybrid retrieval: BM25 (lexical, ru) + FAISS (dense, multilingual-e5)
- Reciprocal Rank Fusion (RRF) merge
- Cross-encoder re-ranking (Jina multilingual reranker)
- Open-source LLM generation (Qwen2.5-7B-Instruct via Ollama or Transformers)
- Deterministic, citation-first JSON answers with doc + clause/table locators

Run (local minimal):
  pip install python-docx razdel rank-bm25 sentence-transformers faiss-cpu transformers torch
  # Optional reranker (recommended): sentence-transformers (already installed)
  # Optional HF 4-bit on Colab: bitsandbytes accelerate

Examples:
  python RAG_prototype_wordaware.py --data_dir ./docs --questions_file ./questions.txt --backend ollama --model Qwen2.5:7b
  python RAG_prototype_wordaware.py --data_dir ./docs --backend hf --model Qwen/Qwen2.5-7B-Instruct --device cuda
"""
from __future__ import annotations
import argparse
import json

import re

import time
from dataclasses import dataclass
from pathlib import Path
from typing import List, Dict, Any, Tuple, Optional

from docx import Document
from razdel import sentenize, tokenize
from rank_bm25 import BM25Okapi
import numpy as np

from sentence_transformers import SentenceTransformer, CrossEncoder
import faiss

# -----------------------------
# Data models
# -----------------------------

@dataclass
class Chunk:
    id: int
    doc_id: int
    doc_name: str
    text: str
    kind: str  # "para" | "table"
    locator: Dict[str, Optional[str]]  # {"clause": str|None, "table": str|None}

    def short_citation(self) -> str:
        parts = [self.doc_name]
        if self.locator.get("clause"):
            parts.append(self.locator["clause"])
        if self.locator.get("table"):
            parts.append(self.locator["table"])
        return ", ".join(parts)

# -----------------------------
# Utilities
# -----------------------------

CLAUSE_RE = re.compile(r"(?:п\.|пункт)\s*\d+(?:\.\d+)*", re.IGNORECASE)
TABLE_RE = re.compile(r"(?:таблица|табл\.)\s*\d+(?:\.\d+)*", re.IGNORECASE)
MULTISPACE_RE = re.compile(r"[ \t]+")


def normalize_text(raw: str) -> str:
    # Keep digits & punctuation. Only collapse whitespace and trim.
    t = MULTISPACE_RE.sub(" ", raw)
    t = re.sub(r"\n{2,}", "\n", t)
    return t.strip()


def extract_locator(text: str) -> Dict[str, Optional[str]]:
    clause = None
    table = None
    m1 = CLAUSE_RE.search(text)
    if m1:
        clause = m1.group(0)
    m2 = TABLE_RE.search(text)
    if m2:
        table = m2.group(0)
    return {"clause": clause, "table": table}

# -----------------------------
# DOCX loader (paragraphs + tables)
# -----------------------------

class DocxLoader:
    def __init__(self, data_dir: Path):
        self.data_dir = Path(data_dir)

    def load(self) -> Tuple[List[Chunk], List[str]]:
        chunks: List[Chunk] = []
        doc_names: List[str] = []
        cid = 0
        for di, path in enumerate(sorted(self.data_dir.glob("*.docx"))):
            doc = Document(str(path))
            shortname = self._shortname_from_filename(path.name)
            doc_names.append(shortname)

            # paragraphs → sentence-aware chunking
            para_texts = [normalize_text(p.text) for p in doc.paragraphs if p.text and p.text.strip()]
            for ch_text in self._word_aware_chunk("\n".join(para_texts)):
                loc = extract_locator(ch_text)
                chunks.append(Chunk(id=cid, doc_id=di, doc_name=shortname, text=ch_text, kind="para", locator=loc))
                cid += 1

            # tables → row-level chunks (attach header as context)
            for tbl in doc.tables:
                header = None
                if len(tbl.rows) > 0:
                    header = " | ".join(cell.text.strip() for cell in tbl.rows[0].cells)
                    header = normalize_text(header)
                for ri, row in enumerate(tbl.rows):
                    if ri == 0:
                        # also create a header chunk (helps retrieval)
                        if header and header.strip():
                            loc = extract_locator(header)
                            chunks.append(Chunk(id=cid, doc_id=di, doc_name=shortname, text=f"[Таблица]\n{header}", kind="table", locator=loc))
                            cid += 1
                        continue
                    row_text = " | ".join(cell.text.strip() for cell in row.cells)
                    row_text = normalize_text(row_text)
                    combined = (f"[Таблица]\n{header}\n{row_text}" if header else f"[Таблица]\n{row_text}")
                    loc = extract_locator(combined)
                    chunks.append(Chunk(id=cid, doc_id=di, doc_name=shortname, text=combined, kind="table", locator=loc))
                    cid += 1
        return chunks, doc_names

    @staticmethod
    def _shortname_from_filename(fname: str) -> str:
        # e.g. "СП 20.13330.2016 Нагрузки...docx" → "СП 20.13330.2016"
        base = Path(fname).stem
        m = re.match(r"^([^\s]+\s*[^\s]+)\b", base)
        return m.group(1) if m else base

    @staticmethod
    def _word_aware_chunk(text: str, max_chars: int = 1600, overlap_sents: int = 1) -> List[str]:
        """Sentence-based sliding window, keeping word boundaries.
        - Split by sentences (razdel)
        - Accumulate into windows up to max_chars
        - Overlap by N sentences to preserve context
        """
        sents = [s.text.strip() for s in sentenize(text) if s.text and s.text.strip()]
        res = []
        i = 0
        while i < len(sents):
            window = []
            total = 0
            j = i
            while j < len(sents) and total + len(sents[j]) + 1 <= max_chars:
                window.append(sents[j])
                total += len(sents[j]) + 1
                j += 1
            if window:
                res.append(" ".join(window))
            if j == i:  # single very long sentence, hard-cut at word boundary
                toks = [t.text for t in tokenize(sents[i])]
                buf = []
                tlen = 0
                for tok in toks:
                    if tlen + len(tok) + 1 > max_chars:
                        break
                    buf.append(tok)
                    tlen += len(tok) + 1
                if buf:
                    res.append(" ".join(buf))
                j = i + 1
            i = max(j - overlap_sents, j)
        return res

# -----------------------------
# Indexes: BM25 + FAISS + RRF
# -----------------------------

class HybridIndex:
    def __init__(self, chunks: List[Chunk], embed_model: str = "intfloat/multilingual-e5-large-instruct", faiss_path: Optional[Path] = None):
        self.chunks = chunks
        self.embed_model_name = embed_model
        self.faiss_path = Path(faiss_path) if faiss_path else None

        # BM25
        tokenized = [[t.text for t in tokenize(c.text.lower())] for c in chunks]
        self._bm25 = BM25Okapi(tokenized)

        # Embeddings
        self._embedder = SentenceTransformer(self.embed_model_name)
        self._embedder_max = 4096
        vectors = self._encode_passages([c.text for c in chunks])
        self.dim = vectors.shape[1]

        # FAISS (cosine via inner product on normalized vectors)
        self._faiss = faiss.IndexFlatIP(self.dim)
        self._faiss.add(vectors)
        if self.faiss_path:
            faiss.write_index(self._faiss, str(self.faiss_path))

    def _encode_passages(self, texts: List[str]) -> np.ndarray:
        batch = []
        embs = []
        for t in texts:
            # e5-instruct expects "passage: " prefix
            batch.append(f"passage: {t[:self._embedder_max]}")
            if len(batch) == 64:
                e = self._embedder.encode(batch, normalize_embeddings=True, convert_to_numpy=True, show_progress_bar=False)
                embs.append(e)
                batch = []
        if batch:
            e = self._embedder.encode(batch, normalize_embeddings=True, convert_to_numpy=True, show_progress_bar=False)
            embs.append(e)
        return np.vstack(embs)

    def _encode_query(self, q: str) -> np.ndarray:
        return self._embedder.encode([f"query: {q}"], normalize_embeddings=True, convert_to_numpy=True, show_progress_bar=False)[0]

    def search(self, query: str, k_bm25: int = 30, k_faiss: int = 30, k_out: int = 25) -> List[Chunk]:
        # BM25
        bm25_scores = self._bm25.get_scores([t.text for t in tokenize(query.lower())])
        top_bm_idx = np.argsort(bm25_scores)[::-1][:k_bm25]

        # FAISS
        qv = self._encode_query(query)
        D, I = self._faiss.search(qv.reshape(1, -1), k_faiss)
        faiss_scores = D[0]
        faiss_idx = I[0]

        # RRF merge
        runs = [list(top_bm_idx), list(faiss_idx)]
        rrf_scores: Dict[int, float] = {}
        K = 60
        for run in runs:
            for rank, idx in enumerate(run, 1):
                rrf_scores[idx] = rrf_scores.get(idx, 0.0) + 1.0 / (K + rank)
        merged = sorted(rrf_scores.items(), key=lambda x: x[1], reverse=True)[:k_out]
        return [self.chunks[i] for i, _ in merged]

# -----------------------------
# Re-ranker (cross-encoder)
# -----------------------------

class Reranker:
    def __init__(self, model: str = "jinaai/jina-reranker-v2-base-multilingual", device: Optional[str] = None):
        # device: "cuda" | "cpu" | None (auto)
        self._ce = CrossEncoder(model, device=device)

    def rerank(self, query: str, docs: List[Chunk], top_k: int = 10) -> List[Chunk]:
        pairs = [[query, d.text] for d in docs]
        scores = self._ce.predict(pairs, batch_size=32, show_progress_bar=False)
        order = np.argsort(scores)[::-1][:top_k]
        return [docs[i] for i in order]

# -----------------------------
# LLM backends
# -----------------------------

class LLMBackend:
    def generate(self, system: str, prompt: str, max_new_tokens: int = 128) -> str:
        raise NotImplementedError

class OllamaBackend(LLMBackend):
    def __init__(self, model: str = "Qwen2.5:7b", host: str = "http://localhost:11434"):
        self.model = model
        self.host = host.rstrip("/")
        try:
            import requests  # lazy
            self._requests = requests
        except Exception as e:
            raise RuntimeError("requests is required for Ollama backend") from e

    def generate(self, system: str, prompt: str, max_new_tokens: int = 128) -> str:
        payload = {
            "model": self.model,
            "stream": False,
            "options": {"temperature": 0, "num_predict": max_new_tokens},
            "system": system,
            "prompt": prompt,
        }
        r = self._requests.post(f"{self.host}/api/generate", json=payload, timeout=600)
        r.raise_for_status()
        data = r.json()
        return data.get("response", "")

class HFLocalBackend(LLMBackend):
    def __init__(self, model: str = "Qwen/Qwen2.5-7B-Instruct", device: Optional[str] = None, load_4bit: bool = False):
        from transformers import AutoModelForCausalLM, AutoTokenizer
        import torch
        self.torch = torch
        self.tok = AutoTokenizer.from_pretrained(model, trust_remote_code=True)
        if load_4bit:
            from transformers import BitsAndBytesConfig
            bnb = BitsAndBytesConfig(load_in_4bit=True, bnb_4bit_compute_dtype=torch.bfloat16)
            self.model = AutoModelForCausalLM.from_pretrained(model, device_map="auto", quantization_config=bnb, torch_dtype=torch.bfloat16, trust_remote_code=True)
        else:
            self.model = AutoModelForCausalLM.from_pretrained(model, device_map="auto", torch_dtype=torch.bfloat16, trust_remote_code=True)

    def generate(self, system: str, prompt: str, max_new_tokens: int = 128) -> str:
        from transformers import TextStreamer
        import torch
        messages = [
            {"role": "system", "content": system},
            {"role": "user", "content": prompt},
        ]
        # Simple chat template handling
        if hasattr(self.tok, "apply_chat_template"):
            text = self.tok.apply_chat_template(messages, add_generation_prompt=True, tokenize=False)
        else:
            text = system + "\n" + prompt
        inputs = self.tok(text, return_tensors="pt").to(self.model.device)
        with torch.no_grad():
            out = self.model.generate(**inputs, do_sample=False, max_new_tokens=max_new_tokens, temperature=0.0)
        gen = self.tok.decode(out[0][inputs["input_ids"].shape[1]:], skip_special_tokens=True)
        return gen

# -----------------------------
# Orchestrator
# -----------------------------

ANSWER_SYSTEM_PROMPT = (
    "Ты — строгий помощник по нормативной документации. Отвечай только по приведённому контексту. "
    "Формат ответа — JSON с полями: {\"answer\": str, \"used_chunks\": [int, ...]}. "
    "Требования: 1) ответ кратко (1–2 предложения), 2) по-русски, 3) при наличии укажи номера пунктов/таблиц в самом ответе, "
    "4) если данных недостаточно — верни 'Недостаточно информации'."
)

ANSWER_USER_PROMPT_TEMPLATE = (
    "Вопрос: {question}\n\nКонтекст (список чанков):\n{context}\n\n"
    "Верни строго JSON без лишнего текста."
)

def build_context(chunks: List[Chunk]) -> str:
    lines = []
    for c in chunks:
        head = f"[ID={c.id}] {c.short_citation()}"
        body = c.text
        lines.append(f"{head}\n{body}\n")
    return "\n".join(lines)

class RAGEngine:
    def __init__(self, data_dir: Path, backend: str = "ollama", model: str = "Qwen2.5:7b", device: Optional[str] = None):
        self.data_dir = Path(data_dir)
        self.backend_kind = backend
        self.model_name = model
        self.device = device

        loader = DocxLoader(self.data_dir)
        self.chunks, self.doc_names = loader.load()
        self.index = HybridIndex(self.chunks)
        self.reranker = Reranker(device=device)
        if backend == "ollama":
            self.llm: LLMBackend = OllamaBackend(model=model)
        elif backend == "hf":
            self.llm = HFLocalBackend(model=model, load_4bit=(device == "cuda"))
        else:
            raise ValueError("backend must be 'ollama' or 'hf'")

    def answer_one(self, question: str, top_k_ctx: int = 8) -> Dict[str, Any]:
        cands = self.index.search(question, k_bm25=30, k_faiss=30, k_out=25)
        top = self.reranker.rerank(question, cands, top_k=top_k_ctx)
        context = build_context(top)
        prompt = ANSWER_USER_PROMPT_TEMPLATE.format(question=question, context=context)
        raw = self.llm.generate(ANSWER_SYSTEM_PROMPT, prompt, max_new_tokens=200)
        # Try to parse JSON
        used_ids: List[int] = []
        answer_text = ""
        try:
            obj = json.loads(extract_json(raw))
            answer_text = obj.get("answer", "").strip()
            used_ids = [int(x) for x in obj.get("used_chunks", [])]
        except Exception:
            # Fallback: use top-3 IDs
            answer_text = raw.strip()
            used_ids = [c.id for c in top[:3]]
        citations = unique_citations(self.chunks, used_ids)
        return {"question": question, "answer": answer_text, "citations": citations}

# -----------------------------
# Helpers
# -----------------------------

def extract_json(s: str) -> str:
    """Extract the first JSON object from a string"""
    start = s.find("{")
    end = s.rfind("}")
    if start != -1 and end != -1 and end > start:
        return s[start:end+1]
    return s


def unique_citations(chunks: List[Chunk], ids: List[int]) -> List[Dict[str, str]]:
    seen = set()
    out = []
    id2chunk = {c.id: c for c in chunks}
    for cid in ids:
        c = id2chunk.get(cid)
        if not c:
            continue
        key = (c.doc_name, c.locator.get("clause"), c.locator.get("table"))
        if key in seen:
            continue
        seen.add(key)
        item = {"document": c.doc_name}
        if c.locator.get("clause"):
            item["clause"] = c.locator["clause"]
        if c.locator.get("table"):
            item["table"] = c.locator["table"]
        out.append(item)
    return out

# -----------------------------
# CLI
# -----------------------------

def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--data_dir", type=str, required=True, help="Directory with .docx files")
    ap.add_argument("--backend", type=str, default="ollama", choices=["ollama", "hf"], help="LLM backend")
    ap.add_argument("--model", type=str, default="Qwen2.5:7b", help="Model name (e.g., 'Qwen2.5:7b' for Ollama or HF repo for --backend hf)")
    ap.add_argument("--device", type=str, default=None, help="Device for reranker/HF ('cuda' or 'cpu')")
    ap.add_argument("--questions_file", type=str, default=None, help="Path to file with questions (one per line)")
    args = ap.parse_args()

    engine = RAGEngine(Path(args.data_dir), backend=args.backend, model=args.model, device=args.device)

    if args.questions_file and Path(args.questions_file).exists():
        questions = [q.strip() for q in Path(args.questions_file).read_text(encoding='utf-8').splitlines() if q.strip()]
    else:
        # Default: the 10 questions from the assignment
        questions = [
            "В каких зонах по весу снежного покрова находятся Херсон и Мелитополь?",
            "Какие регионы Российской Федерации имеют высотный коэффициент k_h, превышающий 2?",
            "Выведи рекомендуемые варианты конструктивного решения заземлителей для стержневых молниеприемников.",
            "Что означает аббревиатура 'ТС'?",
            "Что должна содержать Пояснительная записка в графической части?",
            "Сколько разделов должна содержать проектная документация согласно 87ому постановлению?",
            "Какая максимальная скорость движения подземных машин в выработках?",
            "Какая максимальная температура допускается в горных выработках?",
            "Какие допустимые значения по отклонению геометрических параметров сечения горных выработок?",
            "В каком пункте указана минимальная толщина защитного слоя бетона для арматуры при креплении стволов монолитной бетонной крепью?",
        ]

    results = []
    for q in questions:
        t0 = time.time()
        out = engine.answer_one(q)
        out["latency_sec"] = round(time.time() - t0, 3)
        results.append(out)
        print("\nQ:", q)
        print("A:", out["answer"])
        print("Источники:")
        for c in out["citations"]:
            print(" - ", ", ".join([f"{k}: {v}" for k, v in c.items()]))

    ts = time.strftime("%Y%m%d_%H%M%S")
    out_path = Path(f"results_{ts}.json")
    Path(out_path).write_text(json.dumps(results, ensure_ascii=False, indent=2), encoding='utf-8')
    print(f"\nSaved to {out_path}")


if __name__ == "__main__":
    main()
